import os
import io
import re
import sys
import json
import datetime
import logging
import subprocess
import urllib.request
import urllib.parse
import urllib.error
import ssl
from typing import Optional, Dict, Any, List
from pathlib import Path

logger = logging.getLogger(__name__)

HAS_DOCX = False
Document = None
Inches = None
Pt = None
RGBColor = None
WD_ALIGN_PARAGRAPH = None
WD_TABLE_ALIGNMENT = None
parse_xml = None
nsdecls = None

def ensure_docx_installed() -> bool:
    global HAS_DOCX, Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, parse_xml, nsdecls
    if HAS_DOCX:
        return True
    try:
        import docx
        from docx import Document as _Document
        from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml as _parse_xml
        from docx.oxml.ns import nsdecls as _nsdecls

        Document = _Document
        Inches = _Inches
        Pt = _Pt
        RGBColor = _RGBColor
        WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
        WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
        parse_xml = _parse_xml
        nsdecls = _nsdecls

        HAS_DOCX = True
        return True
    except ImportError:
        logger.info("python-docx missing. Attempting runtime auto-installation via pip...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            import docx
            from docx import Document as _Document
            from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
            from docx.oxml import parse_xml as _parse_xml
            from docx.oxml.ns import nsdecls as _nsdecls

            Document = _Document
            Inches = _Inches
            Pt = _Pt
            RGBColor = _RGBColor
            WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
            WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
            parse_xml = _parse_xml
            nsdecls = _nsdecls

            HAS_DOCX = True
            logger.info("python-docx successfully installed at runtime!")
            return True
        except Exception as e:
            logger.error(f"Runtime auto-install of python-docx failed: {e}")
            return False

def replace_text_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """Replaces text in a paragraph, handling placeholders split across multiple XML runs."""
    if old_text not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    full_text = paragraph.text.replace(old_text, new_text)
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return True
    return False

class TechSpecGenerator:
    """Generates professional AI-powered SAP CPI Technical Specification Word (.docx) documents."""

    NAVY_HEX = "0B2545"
    BLUE_HEX = "134074"
    LIGHT_BG_HEX = "F4F6F9"
    CODE_BG_HEX = "07101D"

    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY") or os.getenv("PALM_API_KEY") or os.getenv("GOOGLE_API_KEY")

    def generate_tech_spec(
        self,
        analysis_data: Dict[str, Any],
        metadata: Optional[Dict[str, Any]] = None,
        reference_docx_bytes: Optional[bytes] = None
    ) -> bytes:
        if not ensure_docx_installed():
            raise RuntimeError("python-docx package is not available and could not be installed automatically.")

        iflow_name = analysis_data.get("name") or (metadata and metadata.get("name")) or "SAP iFlow"
        iflow_id = (metadata and metadata.get("id")) or iflow_name

        # Extract reference document text if available
        reference_text = ""
        if reference_docx_bytes and len(reference_docx_bytes) > 200:
            try:
                ref_doc = Document(io.BytesIO(reference_docx_bytes))
                ref_paras = [p.text for p in ref_doc.paragraphs if p.text.strip()]
                reference_text = "\n".join(ref_paras[:40])
            except Exception as ex_ref:
                logger.warning(f"Could not extract reference text: {ex_ref}")

        # Synthesize AI / Rule-based technical prose
        ai_content = self._synthesize_technical_prose(analysis_data, metadata, reference_text)

        if reference_docx_bytes and len(reference_docx_bytes) > 200:
            try:
                return self._overhaul_fill_reference_doc(reference_docx_bytes, analysis_data, metadata, ai_content)
            except Exception as e:
                logger.error(f"Failed to fill reference docx template: {e}. Falling back to standard docx synthesis.", exc_info=True)

        return self._build_standard_tech_spec(analysis_data, metadata, iflow_name, iflow_id, ai_content)

    def _synthesize_technical_prose(
        self,
        analysis: Dict[str, Any],
        metadata: Optional[Dict[str, Any]],
        reference_text: str = ""
    ) -> Dict[str, Any]:
        iflow_name = analysis.get("name") or (metadata and metadata.get("name")) or "iFlow"
        iflow_id = (metadata and metadata.get("id")) or iflow_name
        sender = analysis.get("sender") or "HTTPS Sender Adapter"
        receiver = analysis.get("receiver") or "Target Receiver System"
        steps_list = analysis.get("steps", [])
        steps_str = ", ".join(steps_list) if steps_list else "HTTPS Adapter, Content Modifier, Message Mapping, Request-Reply"
        inventory = analysis.get("inventory", {})
        groovy_scripts = inventory.get("Groovy Scripts", [])
        schemas = inventory.get("XML Schemas / WSDLs", [])

        # Detailed algorithmic synthesis
        summary_prose = (
            f"This technical specification provides the complete end-to-end interface design, "
            f"BPMN process sequence, extracted configuration parameters, mandatory HTTP headers, "
            f"and schema-derived test payloads for the SAP Integration Suite iFlow '{iflow_name}' (Technical ID: {iflow_id}). "
            f"The interface processes inbound payload triggers from '{sender}' and delivers validated, transformed message payloads to '{receiver}'."
        )

        architecture_prose = (
            f"The interface implements an enterprise message processing flow consisting of {len(steps_list) or 4} key processing steps: "
            f"{steps_str}. "
            f"Inbound HTTP/REST requests are validated against defined XML/JSON schemas. "
            f"Groovy script transformations ({', '.join(groovy_scripts) if groovy_scripts else 'Standard mapping scripts'}) "
            f"execute message modification and dynamic header population prior to target dispatch."
        )

        error_prose = (
            f"Runtime processing failures trigger the iFlow Exception Subprocess. "
            f"The integration engine captures the fault traceback, populates an SAP Message Processing Log ID (MPL ID) "
            f"header ('sap_messageprocessinglogid'), and logs error details for end-to-end monitoring and support."
        )

        default_content = {
            "summary": summary_prose,
            "architecture": architecture_prose,
            "error_handling": error_prose
        }

        if not self.api_key:
            return default_content

        # Call Gemini AI if API key is provided
        try:
            prompt = f"""
You are a Lead SAP CPI Integration Architect writing an official Technical Specification Document.
Generate detailed technical documentation sections for the following SAP CPI iFlow:

iFlow Name: {iflow_name} (ID: {iflow_id})
Sender System: {sender}
Receiver Systems: {receiver}
Flow Sequence: {steps_str}
Groovy Scripts: {', '.join(groovy_scripts) if groovy_scripts else 'None'}
Schemas/WSDLs: {', '.join(schemas) if schemas else 'None'}

Return a valid JSON object containing:
{{
  "summary": "Rich 3-sentence executive summary.",
  "architecture": "Detailed technical architectural description of transformation logic and routing.",
  "error_handling": "Detailed description of error handling, SAP MPL ID tracking, and exception subprocess."
}}
"""
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={self.api_key}"
            payload = json.dumps({"contents": [{"parts": [{"text": prompt}]}]}).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"}, method="POST")

            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE

            with urllib.request.urlopen(req, context=ctx, timeout=12) as resp:
                resp_data = json.loads(resp.read().decode("utf-8"))
                text_out = resp_data["candidates"][0]["content"]["parts"][0]["text"]
                json_match = re.search(r"\{.*\}", text_out, re.DOTALL)
                if json_match:
                    ai_json = json.loads(json_match.group(0))
                    return {
                        "summary": ai_json.get("summary", summary_prose),
                        "architecture": ai_json.get("architecture", architecture_prose),
                        "error_handling": ai_json.get("error_handling", error_prose)
                    }
        except Exception as ex_ai:
            logger.warning(f"Gemini AI synthesis note: {ex_ai}. Using rule-based technical prose.")

        return default_content

    def _set_cell_background(self, cell, fill_hex: str):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def _overhaul_fill_reference_doc(
        self,
        reference_bytes: bytes,
        analysis: Dict[str, Any],
        metadata: Optional[Dict[str, Any]],
        ai_content: Dict[str, Any]
    ) -> bytes:
        """Intelligently updates reference Word document tables, section headings, and placeholders with extracted iFlow data."""
        doc = Document(io.BytesIO(reference_bytes))
        iflow_name = analysis.get("name") or (metadata and metadata.get("name")) or "iFlow"
        iflow_id = (metadata and metadata.get("id")) or iflow_name
        sender = analysis.get("sender") or "HTTPS Sender"
        receiver = analysis.get("receiver") or "Receiver System"
        today_str = datetime.date.today().strftime("%B %d, %Y")

        replacements = {
            "{{IFLOW_NAME}}": iflow_name,
            "{{IFLOW_ID}}": iflow_id,
            "{{SENDER}}": sender,
            "{{RECEIVER}}": receiver,
            "{{DATE}}": today_str,
            "{{SUMMARY}}": ai_content.get("summary", ""),
            "{{ARCHITECTURE}}": ai_content.get("architecture", ""),
            "{{ERROR_HANDLING}}": ai_content.get("error_handling", ""),
            "{{TIMESTAMP}}": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

        # 1. Replace placeholder tags in all paragraphs
        for p in doc.paragraphs:
            for key, val in replacements.items():
                replace_text_in_paragraph(p, key, val)

        # 2. Inspect and replace placeholders inside existing tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in replacements.items():
                            replace_text_in_paragraph(p, key, val)

        # 3. Match and Populate Existing Tables in Reference Document
        configs = analysis.get("config", [])
        headers = analysis.get("headers", [])
        properties = analysis.get("properties", [])

        for table in doc.tables:
            if len(table.rows) == 0:
                continue
            hdr_text = " ".join([c.text.lower() for c in table.rows[0].cells])
            
            # Check if this table is a Configuration Table
            if any(k in hdr_text for k in ["step", "kind", "property", "parameter", "value"]):
                if configs:
                    self._populate_reference_table_rows(table, configs, ["step", "kind", "action", "name", "value"])

            # Check if this table is a Headers Table
            elif "header" in hdr_text or "mandatory" in hdr_text:
                if headers:
                    self._populate_reference_table_rows(table, headers, ["name", "sample", "mandatory", "notes"])

            # Check if this table is an Exchange Properties Table
            elif "property" in hdr_text or "exchange" in hdr_text:
                if properties:
                    self._populate_reference_table_rows(table, properties, ["name", "sample", "mandatory", "notes"])

        # 4. Append Extracted Specifications Section to complete technical spec
        doc.add_heading(f"Extracted Specifications for '{iflow_name}' (ID: {iflow_id})", level=1)
        
        p_intro = doc.add_paragraph()
        p_intro.add_run(ai_content.get("summary", ""))

        doc.add_heading("1. Interface Flow Sequence", level=2)
        steps = analysis.get("steps", [])
        if steps:
            for i, step in enumerate(steps, 1):
                doc.add_paragraph(f"{i}. {step}")
        else:
            doc.add_paragraph(ai_content.get("architecture", ""))

        doc.add_heading("2. Extracted Configuration Table", level=2)
        self._add_config_table(doc, configs)

        doc.add_heading("3. Required HTTP Headers", level=2)
        self._add_requirements_table(doc, "Required HTTP Headers", headers)

        doc.add_heading("4. Required Exchange Properties", level=2)
        self._add_requirements_table(doc, "Exchange Properties", properties)

        doc.add_heading("5. Schema-Derived Test Payloads", level=2)
        self._add_payloads_section(doc, analysis.get("payloads", []))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _populate_reference_table_rows(self, table, rows_data: List[Dict[str, Any]], col_keys: List[str]):
        """Clears old dummy data rows in a reference document table and inserts real extracted iFlow data rows."""
        if len(table.rows) <= 1 or not rows_data:
            return

        # Extract header background fill
        bg_hex = self.NAVY_HEX
        try:
            hdr_cells = table.rows[0].cells
            tcPr = hdr_cells[0]._tc.get_or_add_tcPr()
            shd = tcPr.find(parse_xml(f'<w:shd {nsdecls("w")}/>').tag)
            if shd is not None and shd.attrib.get(f'{{{nsdecls("w")}}}fill'):
                bg_hex = shd.attrib.get(f'{{{nsdecls("w")}}}fill')
        except Exception:
            pass

        # Clear existing data rows except header
        num_rows = len(table.rows)
        for _ in range(num_rows - 1):
            tr = table.rows[1]._tr
            table._tbl.remove(tr)

        # Insert new iFlow data rows
        for rdata in rows_data:
            new_row = table.add_row()
            for idx, key in enumerate(col_keys):
                if idx < len(new_row.cells):
                    val = str(rdata.get(key, ""))
                    if key == "mandatory":
                        val = "Yes" if rdata.get(key) else "No"
                    new_row.cells[idx].text = val

    def _build_standard_tech_spec(
        self,
        analysis: Dict[str, Any],
        metadata: Optional[Dict[str, Any]],
        iflow_name: str,
        iflow_id: str,
        ai_content: Dict[str, Any]
    ) -> bytes:
        doc = Document()

        # Page Margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Title Banner
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run("SAP CPI Technical Specification Document")
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(24)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(11, 37, 69)

        # Subtitle
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(18)
        run_sub = p_sub.add_run(f"Integration Artifact: {iflow_name} (ID: {iflow_id})")
        run_sub.font.name = "Calibri"
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = RGBColor(19, 64, 116)

        # Document Metadata Table
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.style = 'Table Grid'
        
        meta_data = [
            ("Technical Artifact ID", iflow_id),
            ("Integration Name", iflow_name),
            ("Document Version / Date", f"1.0.0 — {datetime.date.today().strftime('%B %d, %Y')}"),
            ("Status", "Approved Technical Specification")
        ]
        for row_idx, (k, v) in enumerate(meta_data):
            r_cells = meta_table.rows[row_idx].cells
            r_cells[0].text = k
            r_cells[1].text = v
            self._set_cell_background(r_cells[0], self.LIGHT_BG_HEX)
            for p in r_cells[0].paragraphs:
                for r in p.runs:
                    r.font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # 1. Executive Summary
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(ai_content.get("summary", ""))

        # Package Inventory
        inventory = analysis.get("inventory", {})
        if inventory:
            p_inv = doc.add_paragraph()
            p_inv.add_run("Package Artifact Inventory:").bold = True
            for kind, files in inventory.items():
                if files:
                    doc.add_paragraph(f"• {kind}: " + ", ".join(files), style='List Bullet')

        # 2. Interface Architecture & Sequence
        doc.add_heading("2. Interface Architecture & Flow Sequence", level=1)
        doc.add_paragraph(ai_content.get("architecture", ""))
        doc.add_paragraph(f"• Sender System: {analysis.get('sender', 'HTTPS Sender Adapter')}")
        doc.add_paragraph(f"• Receiver Systems: {analysis.get('receiver', 'Backend Target System')}")

        steps = analysis.get("steps", [])
        if steps:
            doc.add_paragraph("Execution Steps Sequence:").bold = True
            for i, step in enumerate(steps, 1):
                doc.add_paragraph(f"{i}. {step}")

        # 3. Exception Handling
        doc.add_heading("3. Exception Handling & Logging", level=1)
        doc.add_paragraph(ai_content.get("error_handling", ""))

        # 4. Configuration Table
        doc.add_heading("4. Extracted Configuration Table", level=1)
        self._add_config_table(doc, analysis.get("config", []))

        # 5. Required Headers
        doc.add_heading("5. Required HTTP Headers", level=1)
        self._add_requirements_table(doc, "Required HTTP Headers", analysis.get("headers", []))

        # 6. Exchange Properties
        doc.add_heading("6. Required Exchange Properties", level=1)
        self._add_requirements_table(doc, "Exchange Properties", analysis.get("properties", []))

        # 7. Test Payloads Section
        doc.add_heading("7. Schema-Derived Test Payloads", level=1)
        self._add_payloads_section(doc, analysis.get("payloads", []))

        # 8. Assumptions & Gaps
        assumptions = analysis.get("assumptions", [])
        if assumptions:
            doc.add_heading("8. Assumptions & Technical Gaps", level=1)
            for ass in assumptions:
                doc.add_paragraph(f"• {ass}", style='List Bullet')

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_config_table(self, doc, configs: List[Dict[str, Any]]):
        if not configs:
            doc.add_paragraph("No specific configuration steps extracted.")
            return

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_titles = ["Step", "Type", "Action", "Name", "Value / Source"]
        for i, title in enumerate(hdr_titles):
            hdr_cells[i].text = title
            self._set_cell_background(hdr_cells[i], self.NAVY_HEX)
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for row_data in configs:
            r_cells = table.add_row().cells
            r_cells[0].text = str(row_data.get("step", ""))
            r_cells[1].text = str(row_data.get("kind", ""))
            r_cells[2].text = str(row_data.get("action", ""))
            r_cells[3].text = str(row_data.get("name", ""))
            r_cells[4].text = str(row_data.get("value", ""))

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_requirements_table(self, doc, label: str, reqs: List[Dict[str, Any]]):
        if not reqs:
            doc.add_paragraph(f"No mandatory {label.lower()} detected.")
            return

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_titles = ["Name", "Sample Value", "Mandatory?", "Notes"]
        for i, title in enumerate(hdr_titles):
            hdr_cells[i].text = title
            self._set_cell_background(hdr_cells[i], self.BLUE_HEX)
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for req in reqs:
            r_cells = table.add_row().cells
            r_cells[0].text = str(req.get("name", ""))
            r_cells[1].text = str(req.get("sample", ""))
            r_cells[2].text = "Yes" if req.get("mandatory") else "No"
            r_cells[3].text = str(req.get("notes", ""))

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_payloads_section(self, doc, payloads: List[Dict[str, Any]]):
        if not payloads:
            doc.add_paragraph("No schema test payloads generated.")
            return

        for p in payloads:
            scenario = p.get("scenario", "Test Scenario")
            fmt = str(p.get("format", "xml")).upper()
            body = p.get("body", "")
            source = p.get("source", "Schema")

            h3 = doc.add_heading(f"Scenario: {scenario} ({fmt})", level=2)
            doc.add_paragraph(f"Derived from schema: `{source}`")

            p_code = doc.add_paragraph()
            p_code.paragraph_format.space_before = Pt(4)
            p_code.paragraph_format.space_after = Pt(12)
            run_code = p_code.add_run(body)
            run_code.font.name = "Consolas"
            run_code.font.size = Pt(9.5)
